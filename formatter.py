from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

def set_cell_rtl(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    bidi = OxmlElement('w:bidi')
    tcPr.append(bidi)

def insert_images_ai_style(doc_path, images_folder, output_path):
    doc = Document(doc_path)
    section = doc.sections[0]
    section.right_margin = Inches(1)
    section.left_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # הגדרת כיוון כתיבה מימין לשמאל למסמך
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.right_to_left = True

    images = sorted(os.listdir(images_folder))  # אין הגבלה לכמות התמונות
    img_idx = 0
    total_images = len(images)
    page_number = 1

    while img_idx < total_images:
        table = doc.add_table(rows=2, cols=2)
        table.autofit = True

        image_number = 1

        for row in range(2):
            for col in range(2):
                if img_idx >= total_images:
                    break
                img_path = os.path.join(images_folder, images[img_idx])
                cell = table.cell(row, col)

                paragraph = cell.paragraphs[0]
                paragraph.paragraph_format.right_to_left = True
                run = paragraph.add_run()
                run.add_picture(img_path, width=Inches(2.5))
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # מספר תמונה
                caption = cell.add_paragraph(f"{page_number}.{image_number}")
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption.runs[0].font.size = Pt(10)
                caption.paragraph_format.right_to_left = True

                img_idx += 1
                image_number += 1

        doc.add_paragraph("\n")

        desc_table = doc.add_table(rows=1, cols=1)
        desc_cell = desc_table.cell(0, 0)
        desc_cell.height = Inches(1.0)

        # הגדרת מסגרת שחורה לתא עם שמות תקניים
        tcBorders = OxmlElement("w:tcBorders")
        for side in ['top', 'bottom', 'left', 'right']:
            border = OxmlElement(f"w:{side}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "4")
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), "000000")
            tcBorders.append(border)
        desc_cell._tc.get_or_add_tcPr().append(tcBorders)

        # יישור התא והוספת כיתוב
        set_cell_rtl(desc_cell)
        desc_paragraph = desc_cell.paragraphs[0]
        desc_paragraph.paragraph_format.right_to_left = True
        desc_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        desc_paragraph.add_run("תיאור תמונות:")

        doc.add_paragraph("")
        page_number += 1

        if img_idx < total_images:
            doc.add_page_break()

    doc.save(output_path)
    return output_path
