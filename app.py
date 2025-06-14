import os
import shutil
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import streamlit as st

# הגדרת תצוגת העמוד - חובה לפני כל קריאה ל-st
st.set_page_config(page_title="Word Report Generator", layout="centered")
from streamlit_cropper import st_cropper
from streamlit_drawable_canvas import st_canvas

# פונקציית עזר לעריכת תמונה

def edit_image_workflow(img_file, index):
    st.subheader(f"\u05e2\u05e8\u05d9\u05db\u05ea \u05ea\u05de\u05d5\u05e0\u05d4 {index+1}: {img_file.name}")
    image = Image.open(img_file)

    # חיתוך
    cropped_img = st_cropper(image, box_color='#FA8072', aspect_ratio=None)
    st.image(cropped_img, caption="\u05d0\u05d6\u05d5\u05e8 \u05e9\u05e0\u05d1\u05d7\u05e8 \u05dc\u05d7\u05d9\u05ea\u05d5\u05da")

    # ציור
    canvas_result = st_canvas(
        fill_color="rgba(255, 255, 0, 0.3)",
        stroke_width=3,
        stroke_color="#000000",
        background_image=cropped_img,
        update_streamlit=True,
        height=cropped_img.height,
        width=cropped_img.width,
        drawing_mode="freedraw",
        key=f"canvas_{index}",
    )

    if canvas_result.image_data is not None:
        edited_img = Image.fromarray(canvas_result.image_data.astype("uint8"))
        output_path = os.path.join("input/images", f"edited_{img_file.name}")
        edited_img.save(output_path, format="PNG")
        st.success(f"\u05e0\u05e9\u05de\u05e8 \u05e2\u05dd \u05d4\u05e1\u05d9\u05de\u05d5\u05e0\u05d9\u05dd: {output_path}")
        return output_path
    return None

# פונקציית שיבוץ תמונות למסמך

def insert_images_ai_style(doc_path, images_folder, output_path):
    doc = Document(doc_path)
    section = doc.sections[0]
    section.right_margin = Inches(1)
    section.left_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.right_to_left = True

    images = sorted(os.listdir(images_folder))
    img_idx = 0
    total_images = len(images)
    page_number = 1

    while img_idx < total_images:
        table = doc.add_table(rows=2, cols=2)
        table.autofit = True

        image_number = 1

        for row in range(2):
            for col in range(2):
                if img_idx >= total_images:
                    break
                img_path = os.path.join(images_folder, images[img_idx])
                cell = table.cell(row, col)
                paragraph = cell.paragraphs[0]
                paragraph.paragraph_format.right_to_left = True
                run = paragraph.add_run()
                run.add_picture(img_path, width=Inches(2.5))
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                caption = cell.add_paragraph(f"{page_number}.{image_number}")
                caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                caption.runs[0].font.size = Pt(10)
                caption.paragraph_format.right_to_left = True

                img_idx += 1
                image_number += 1

        doc.add_paragraph("\n")
        desc_table = doc.add_table(rows=1, cols=1)
        desc_cell = desc_table.cell(0, 0)
        desc_cell.height = Inches(1.0)

        tcBorders = OxmlElement("w:tcBorders")
        for side in ['top', 'bottom', 'left', 'right']:
            border = OxmlElement(f"w:{side}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "4")
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), "000000")
            tcBorders.append(border)
        desc_cell._tc.get_or_add_tcPr().append(tcBorders)

        set_cell_rtl(desc_cell)
        desc_paragraph = desc_cell.paragraphs[0]
        desc_paragraph.paragraph_format.right_to_left = True
        desc_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        desc_paragraph.add_run("\u05ea\u05d9\u05d0\u05d5\u05e8 \u05ea\u05de\u05d5\u05e0\u05d5\u05ea:")

        doc.add_paragraph("")
        page_number += 1

        if img_idx < total_images:
            doc.add_page_break()

    doc.save(output_path)
    return output_path

# יישור תא RTL

def set_cell_rtl(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    tcPr.append(bidi)

# CSS RTL
st.markdown("""
    <style>
        .stTextInput > div > div > input {
            direction: rtl;
            text-align: right;
        }
        .stTextArea > div > div > textarea {
            direction: rtl;
            text-align: right;
        }
        .stMarkdown {
            direction: rtl;
            text-align: right;
        }
    </style>
""", unsafe_allow_html=True)

st.set_page_config(page_title="Word Report Generator", layout="centered")
st.markdown("## 📄 מחולל דוחות Word עם תמונות (כולל עריכה)", unsafe_allow_html=True)

if 'restart' not in st.session_state:
    st.session_state.restart = False

uploaded_file = st.file_uploader("העלה קובץ Word (.docx בלבד)", type=["docx"])
uploaded_images = st.file_uploader("העלה תמונות", type=["jpg", "jpeg", "png"], accept_multiple_files=True)

if uploaded_file and uploaded_file.name.endswith(".doc"):
    st.error("⚠️ לא ניתן להעלות קבצי .doc ישנים. שמור את הקובץ כ־.docx ונסה שוב.")
    st.stop()

if st.button("\u05e4\u05e8\u05e1 \u05d3\u05d5\u05d7"):
    if not uploaded_file or not uploaded_images:
        st.error("\u05d9\u05e9 \u05dc\u05d4\u05e2\u05dc\u05d5\u05ea \u05d2\u05dd \u05e7\u05d5\u05d1\u05e5 Word \u05d5\u05d2\u05dd \u05ea\u05de\u05d5\u05e0\u05d5\u05ea")
    else:
        try:
            if os.path.exists("input"):
                shutil.rmtree("input")
            if os.path.exists("output"):
                shutil.rmtree("output")
            os.makedirs("input/images", exist_ok=True)
            os.makedirs("output", exist_ok=True)

            input_path = os.path.join("input", uploaded_file.name)
            with open(input_path, "wb") as f:
                f.write(uploaded_file.getbuffer())

            final_image_paths = []
            for idx, img in enumerate(uploaded_images):
                st.write(f"---\n### תמונה {idx+1}: {img.name}")
                edit_mode = st.checkbox(f"⬇️ ערוך את {img.name}", key=f"edit_{idx}")

                if edit_mode:
                    result = edit_image_workflow(img, idx)
                    if result:
                        final_image_paths.append(result)
                else:
                    raw_path = os.path.join("input/images", f"original_{img.name}")
                    with open(raw_path, "wb") as f:
                        f.write(img.getbuffer())
                    final_image_paths.append(raw_path)

            output_path = os.path.join("output", "ready_report.docx")
            final_path = insert_images_ai_style(input_path, "input/images", output_path)

            with open(final_path, "rb") as f:
                report_data = f.read()

            st.success("\ud83d\udcc4 הדו\"ח הופק בהצלחה!")
            st.download_button("\u05d4\u05d5\u05e8\u05d3 \u05d0\u05ea \u05d4\u05d3\u05d5\"\u05d7", data=report_data, file_name="ready_report.docx")

            if st.button("\u05d4\u05e4\u05e7 \u05d3\u05d5\u05d7 \u05e0\u05d5\u05e1\u05e3"):
                st.session_state.restart = True

        except Exception as e:
            st.error(f"\u05e9\u05d2\u05d9\u05d0\u05d4: {str(e)}")

if st.session_state.restart:
    st.session_state.restart = False
    st.experimental_rerun()
